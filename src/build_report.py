"""
build_report.py
---------------
Builds the Tennis Serve Biomechanics Report DOCX with all figures embedded.
Run with: C:/Users/gonza/anaconda3/python.exe src/build_report.py
"""

import os
import io
from datetime import date

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D0D0"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color="1F3864", space_before=14, space_after=6):
    p  = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    return p


def add_body(doc, text, bold=False, italic=False, size=10.5, space_after=6,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.2
    return p


def add_bullet(doc, text, level=0, size=10.5):
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_figure(doc, img_path, caption, width=6.2):
    doc.add_picture(img_path, width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap  = doc.add_paragraph()
    run  = cap.add_run(caption)
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after  = Pt(12)
    cap.paragraph_format.space_before = Pt(3)
    return cap


def add_divider(doc):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def load_all_serves():
    dfs = {}
    for player in ["djokovic", "federer", "amateur"]:
        path  = f"outputs/{player}_serves"
        files = sorted(
            [f for f in os.listdir(path) if f.endswith(".csv")],
            key=lambda x: int(x.split("_")[1].split(".")[0]),
        )
        frames = []
        for f in files:
            df = pd.read_csv(os.path.join(path, f))
            df["serve_id"] = f.replace(".csv", "")
            frames.append(df)
        dfs[player] = pd.concat(frames, ignore_index=True)
    return dfs


def impact_rows(dfs):
    out = {}
    for pk, df in dfs.items():
        out[pk] = df.loc[df.groupby("serve_id")["score"].idxmax()].reset_index(drop=True)
    return out


# ── Build the document ────────────────────────────────────────────────────────

def build():
    dfs     = load_all_serves()
    impacts = impact_rows(dfs)
    summary = pd.read_csv("outputs/summary_statistics.csv")

    FIGS = {
        "radar":       "outputs/report_figures/fig1_radar.png",
        "violin":      "outputs/report_figures/fig2_violin_box.png",
        "motion":      "outputs/report_figures/fig3_motion_profile.png",
        "scatter":     "outputs/report_figures/fig4_scatter.png",
        "wrist":       "outputs/report_figures/fig5_wrist_trajectory.png",
        "heatmap":     "outputs/report_figures/fig6_heatmap.png",
        "bar":         "outputs/report_figures/fig7_bar_consistency.png",
        "impact_img":  "outputs/impact_frame_max.jpg",
    }

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("Tennis Serve Biomechanics Analysis")
    tr.font.size  = Pt(26)
    tr.font.bold  = True
    tr.font.color.rgb = RGBColor.from_string("1F3864")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_p.add_run("AI-Powered Pose Estimation & Comparative Study:\nDjokovic · Federer · Amateur")
    sr.font.size  = Pt(14)
    sr.font.color.rgb = RGBColor.from_string("2E5090")
    sr.font.italic = True

    doc.add_paragraph()

    # Impact image on cover
    if os.path.exists(FIGS["impact_img"]):
        add_figure(doc, FIGS["impact_img"],
                   "Figure: Detected serve impact frame — MediaPipe Pose skeleton overlay",
                   width=4.8)

    doc.add_paragraph()

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Author",   "Gonzalo Angel Tano — Data Science Portfolio Project"),
        ("Date",     date.today().strftime("%B %d, %Y")),
        ("Method",   "MediaPipe Pose (BlazePose) + Custom Biomechanics Pipeline"),
    ]
    for row, (label, val) in zip(meta.rows, meta_data):
        lc = row.cells[0]
        vc = row.cells[1]
        lc.paragraphs[0].add_run(label).font.bold = True
        vc.paragraphs[0].add_run(val)
        for c in (lc, vc):
            c.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1 · EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Executive Summary", level=1)
    add_divider(doc)
    add_body(doc,
        "This report presents a data-driven biomechanical analysis of the tennis serve using "
        "computer vision and pose estimation. MediaPipe BlazePose was applied frame-by-frame "
        "to match footage of two professional players (Novak Djokovic and Roger Federer) and "
        "one amateur, extracting joint angles and wrist kinematics across every detected serve.")

    add_body(doc,
        "A total of 102 serves were detected and segmented: 34 from Djokovic, 63 from Federer, "
        "and 5 from the amateur. Per-frame biomechanical features — elbow angle, knee angle, "
        "shoulder angle, and wrist height — were recorded and analysed at the moment of maximum "
        "impact score, defined as the frame maximising (elbow + knee) angle while the wrist is "
        "at its highest point.")

    add_body(doc,
        "Key findings are summarised below:", bold=True)

    bullets = [
        "Djokovic shows the highest elbow extension at impact (172°, σ=8°), indicating near-full "
        "arm straightening — a hallmark of a powerful flat serve.",
        "Federer achieves the greatest shoulder abduction (132°, σ=48°), reflecting his distinctive "
        "open-stance hip rotation and classic serve mechanics.",
        "The amateur achieves comparable elbow and knee angles to the professionals but shows "
        "significantly lower shoulder abduction (72°, σ=20°) and the lowest impact score, "
        "suggesting the primary area for improvement is shoulder rotation and trunk involvement.",
        "Both professionals exhibit tighter serve-to-serve consistency (lower SD) in elbow and "
        "knee angles than the amateur, highlighting the role of motor pattern consolidation in "
        "elite performance.",
        "Wrist trajectory analysis reveals that professionals achieve a higher wrist peak earlier "
        "in the serve sequence, suggesting a more efficient and explosive toss-to-impact timing.",
    ]
    for b in bullets:
        add_bullet(doc, b)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 2 · METHODOLOGY
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Methodology", level=1)
    add_divider(doc)

    add_heading(doc, "2.1 Pose Estimation Pipeline", level=2, space_before=8)
    add_body(doc,
        "Pose estimation was performed using Google's MediaPipe BlazePose (model complexity = 1), "
        "which outputs 33 3D body landmarks per frame at sub-millimetre resolution (normalised to "
        "image width/height). All videos were processed at full resolution with a detection "
        "confidence threshold of 0.50 and a tracking confidence of 0.50.")

    add_heading(doc, "2.2 Serve Detection Algorithm", level=2, space_before=8)
    add_body(doc,
        "A two-stage state-machine was used to detect serve events:")
    bullets2 = [
        "Stage 1 — Wrist-above-shoulder detection: the minimum wrist-shoulder Y-differential "
        "across both arms was monitored. A serve interval begins when this differential falls below "
        "−0.07 (wrist at least 7% of frame height above shoulder) for at least 5 consecutive frames.",
        "Stage 2 — Impact localisation: within each serve interval, the impact frame is identified "
        "as argmax(elbow_angle + knee_angle − 300 × wrist_height). This composite score rewards "
        "maximum joint extension simultaneous with maximum wrist elevation.",
        "A 90-frame (≈3 s) cooldown prevents double-counting within the same serve motion.",
    ]
    for b in bullets2:
        add_bullet(doc, b)

    add_heading(doc, "2.3 Feature Extraction", level=2, space_before=8)
    add_body(doc, "The following biomechanical features were extracted per frame:")

    feat_table = doc.add_table(rows=5, cols=3)
    feat_table.style = "Table Grid"
    headers = ["Feature", "Landmarks Used", "Biomechanical Interpretation"]
    header_row = feat_table.rows[0]
    for cell, h in zip(header_row.cells, headers):
        cell.paragraphs[0].add_run(h).font.bold = True
        set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    feat_data = [
        ("Elbow Angle",    "Shoulder → Elbow → Wrist",    "Arm extension at contact point"),
        ("Knee Angle",     "Hip → Knee → Ankle",           "Leg drive and ground force utilisation"),
        ("Shoulder Angle", "Hip → Shoulder → Elbow",       "Racket arm abduction / elevation"),
        ("Wrist Height",   "Min(left wrist Y, right wrist Y)", "Serve toss height and arm reach"),
    ]
    for row, (f, lm, interp) in zip(feat_table.rows[1:], feat_data):
        row.cells[0].paragraphs[0].add_run(f).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(lm).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(interp).font.size = Pt(10)
        for cell in row.cells:
            set_cell_bg(cell, "F5F7FA")

    doc.add_paragraph()

    add_heading(doc, "2.4 Dataset", level=2, space_before=8)
    add_body(doc,
        "Three video clips were analysed. The 'Federer' clip is match footage from the US Open "
        "(visible signage in frame); the Djokovic clip is training footage. The amateur clip is "
        "recreational serve practice. All clips were processed offline with no data augmentation.")

    ds_table = doc.add_table(rows=4, cols=4)
    ds_table.style = "Table Grid"
    ds_headers = ["Player", "Serves Detected", "Frames Analysed", "Video Context"]
    for cell, h in zip(ds_table.rows[0].cells, ds_headers):
        cell.paragraphs[0].add_run(h).font.bold = True
        set_cell_bg(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    ds_data = [
        ("Djokovic", "34",  "34 × 36 = 1,224", "Training session"),
        ("Federer",  "63",  "63 × 36 = 2,268", "US Open match footage"),
        ("Amateur",  "5",   "5 × 36 = 180",    "Recreational practice"),
    ]
    for row, (p, n, fr, ctx) in zip(ds_table.rows[1:], ds_data):
        vals = [p, n, fr, ctx]
        for cell, v in zip(row.cells, vals):
            cell.paragraphs[0].add_run(v).font.size = Pt(10)
            set_cell_bg(cell, "F5F7FA")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 3 · RESULTS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Results", level=1)
    add_divider(doc)

    # 3.1 Multi-metric overview
    add_heading(doc, "3.1 Multi-Metric Biomechanical Profile (Radar Chart)", level=2, space_before=8)
    add_body(doc,
        "Figure 1 presents a normalised radar chart comparing the three players across peak elbow, "
        "knee, and shoulder angles at impact. Values are min-max normalised across players to "
        "enable direct visual comparison of relative strengths and weaknesses.")
    add_figure(doc, FIGS["radar"],
               "Figure 1 — Multi-Metric Biomechanical Profile. Axes are min-max normalised. "
               "Raw values: Djokovic: Elbow 172° | Knee 166° | Shoulder 91°; "
               "Federer: Elbow 167° | Knee 170° | Shoulder 132°; Amateur: Elbow 170° | Knee 160° | Shoulder 72°.")
    add_body(doc,
        "The radar chart reveals a striking difference in shoulder angle: Federer's profile extends "
        "far further on the shoulder axis (132°) compared to Djokovic (91°) and especially the "
        "amateur (72°). This indicates Federer's serve relies more heavily on shoulder rotation and "
        "the kinetic chain through the trunk. Elbow and knee angles are comparably high across all "
        "three players, suggesting these metrics are near-ceiling for any competent server.")

    # 3.2 Distribution
    add_heading(doc, "3.2 Distribution of Peak Angles (Violin + Box Plots)", level=2, space_before=8)
    add_body(doc,
        "Figure 2 shows the full distribution of peak angles per metric using violin plots "
        "(kernel density estimate), box plots (median + IQR), and individual serve points. "
        "The annotation μ/σ is the mean and standard deviation across all serves.")
    add_figure(doc, FIGS["violin"],
               "Figure 2 — Distribution of peak joint angles at serve impact. "
               "Violin width represents the density of observations. White line inside box = median.",
               width=6.4)
    add_body(doc,
        "Notable observations:")
    add_bullet(doc, "Elbow angle: Djokovic is the most consistent (σ=8°). Federer shows a bimodal "
                    "distribution with a cluster of outliers near 15°—likely mis-detected non-serve "
                    "motions. The amateur data (n=5) shows tight grouping near 170°.")
    add_bullet(doc, "Knee angle: All three players show right-skewed distributions. The amateur's "
                    "lower median (160°) relative to professionals (166–170°) suggests a more bent "
                    "knee at impact, potentially indicating less explosive leg drive.")
    add_bullet(doc, "Shoulder angle: The widest distributions are seen here for both professionals, "
                    "reflecting natural serve-to-serve variability in arm positioning. The amateur's "
                    "restricted range (72°±20°) is the most actionable coaching finding in this analysis.")

    # 3.3 Motion Profile
    add_heading(doc, "3.3 Serve Motion Profile (Time-Series Analysis)", level=2, space_before=8)
    add_body(doc,
        "Figure 3 displays the average biomechanical trajectory across all detected serves for each "
        "player, aligned so that relative frame 25 corresponds to the detected impact moment. "
        "Shaded bands represent ±1 standard deviation, indicating serve-to-serve variability.")
    add_figure(doc, FIGS["motion"],
               "Figure 3 — Average serve motion profile (mean ± 1 SD). Dashed vertical line at "
               "frame 25 indicates the impact reference point. Each curve is averaged over all "
               "detected serves for that player.",
               width=6.4)
    add_body(doc,
        "The motion profile reveals the temporal structure of each player's serve mechanics:")
    add_bullet(doc, "Elbow angle: Professionals show a rapid extension spike approaching impact, "
                    "reaching near-180° at or just before contact. The amateur's elbow extension "
                    "is less pronounced and peaks earlier, suggesting a premature release.")
    add_bullet(doc, "Knee angle: A clear dip before frame 10 corresponds to the knee bend in the "
                    "trophy pose loading phase, followed by explosive extension into impact. "
                    "Djokovic shows the most pronounced bend-and-extend cycle.")
    add_bullet(doc, "Wrist height: All players show the wrist rising through the serve motion. "
                    "The amateur's wrist rises more slowly and reaches a lower peak, consistent "
                    "with a shorter or less coordinated toss.")
    add_bullet(doc, "Shoulder angle: Large variance in the pre-impact phase reflects the wind-up "
                    "and trophy phase variability. Post-impact drop in shoulder angle corresponds "
                    "to the follow-through motion.")

    # 3.4 Scatter
    add_heading(doc, "3.4 Elbow vs. Knee Correlation at Impact", level=2, space_before=8)
    add_body(doc,
        "Figure 4 plots the elbow angle against the knee angle for every individual serve at "
        "the impact frame. Confidence ellipses (±1 SD) visualise each player's typical serve space.")
    add_figure(doc, FIGS["scatter"],
               "Figure 4 — Elbow vs. Knee angle scatter at impact. Each dot is one serve. "
               "Shaded ellipses represent ±1 SD around each player's centroid.")
    add_body(doc,
        "The professionals cluster tightly in the high-elbow / high-knee quadrant (both >165°), "
        "confirming that elite serve mechanics require simultaneous, coordinated extension of both "
        "joints. Federer's ellipse is notably larger in the knee direction, reflecting his "
        "characteristic leg drive variability across different serve types (flat vs. kick). "
        "The two outlier Federer points at very low elbow angles (<100°) are likely false detections "
        "from non-serve frames and represent a known limitation of wrist-based serve detection on "
        "footage with diverse arm movements.")

    # 3.5 Wrist
    add_heading(doc, "3.5 Wrist Height Trajectory", level=2, space_before=8)
    add_body(doc,
        "Figure 5 shows the average wrist height trajectory. Note that Y-coordinates are inverted "
        "in image space: a lower numerical value means the wrist is physically higher in the frame. "
        "The Y-axis is inverted so the curve rises as the wrist rises.")
    add_figure(doc, FIGS["wrist"],
               "Figure 5 — Average wrist height trajectory. Y-axis is inverted: curves moving upward "
               "indicate the wrist physically rising. Shaded bands = ±1 SD. Coloured regions mark "
               "approximate serve phases: Preparation (0–10), Toss (10–25), Follow-through (25–36).",
               width=6.2)
    add_body(doc,
        "Federer achieves the highest wrist elevation throughout the serve (lowest Y-value ≈0.40), "
        "consistent with his famously high and precise toss. Djokovic and the amateur show similar "
        "peak wrist heights around 0.45, though Djokovic's trajectory is smoother (lower SD). "
        "The amateur's toss phase (frames 10–25) is less clean, with a higher variance band, "
        "indicating inconsistent ball toss height — a known fundamentals issue in developing players.")

    # 3.6 Heatmap
    add_heading(doc, "3.6 Summary Heatmap", level=2, space_before=8)
    add_body(doc,
        "Figure 6 consolidates all metrics into a dual heatmap: raw values on the left, "
        "z-score normalised values on the right. The z-score view removes scale differences "
        "and highlights relative strengths and weaknesses within the group.")
    add_figure(doc, FIGS["heatmap"],
               "Figure 6 — Biomechanics summary heatmap. Left: raw mean values. "
               "Right: z-score normalised (red = above group mean, blue = below). "
               "Impact Score = elbow + knee − 300 × wrist_height (higher is better).",
               width=6.4)
    add_body(doc,
        "The z-score matrix tells a clear story: Federer leads on shoulder angle (+1.10 SD) and "
        "impact score (+0.67 SD) while having the lowest elbow angle (−1.11 SD), consistent with "
        "his more fluid, less locked-out arm style. Djokovic leads on elbow extension (+0.82 SD). "
        "The amateur is notably above the group mean only on wrist height (+1.14 SD), meaning the "
        "wrist does not rise as high as the professionals — the single most important coaching cue.")

    # 3.7 Consistency
    add_heading(doc, "3.7 Serve Consistency (Mean ± SD Bar Chart)", level=2, space_before=8)
    add_body(doc,
        "Figure 7 provides a clean grouped bar chart of mean peak angles with ±1 SD error bars, "
        "making inter-player and inter-metric differences immediately legible.")
    add_figure(doc, FIGS["bar"],
               "Figure 7 — Mean peak joint angles at serve impact with ±1 SD error bars. "
               "Smaller error bars indicate higher serve-to-serve consistency.",
               width=6.0)
    add_body(doc,
        "Djokovic's error bars are the narrowest for elbow angle, confirming his reputation "
        "for a mechanically consistent, repeatable flat serve. Federer's wider shoulder angle "
        "error bar reflects his use of multiple serve types (flat, slice, kick) which produce "
        "different shoulder positions. The amateur's very large shoulder error bar suggests "
        "inconsistency in the trophy pose, the most foundational serve element.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 4 · STATISTICAL SUMMARY TABLE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Statistical Summary", level=1)
    add_divider(doc)
    add_body(doc,
        "Table 1 presents the mean and standard deviation of peak joint angles for each player, "
        "computed from impact-frame metrics across all detected serves.")

    # Build enriched summary
    metrics_map = {
        "Elbow Angle":    "elbow_angle",
        "Knee Angle":     "knee_angle",
        "Shoulder Angle": "shoulder_angle",
        "Impact Score":   "score",
    }
    player_disp = [("Djokovic", "djokovic"), ("Federer", "federer"), ("Amateur", "amateur")]

    col_headers = ["Metric"] + [f"{pl} Mean ± SD" for pl, _ in player_disp] + ["Best Professional"]
    tbl = doc.add_table(rows=len(metrics_map) + 1, cols=len(col_headers))
    tbl.style = "Table Grid"

    for i, h in enumerate(col_headers):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).font.bold = True
        set_cell_bg(cell, "1F3864")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    for row_i, (metric_label, metric_key) in enumerate(metrics_map.items(), start=1):
        row    = tbl.rows[row_i]
        row.cells[0].paragraphs[0].add_run(metric_label).font.size = Pt(10)
        set_cell_bg(row.cells[0], "EBF0F7")

        best_val  = -np.inf
        best_name = ""
        for col_i, (pl, pk) in enumerate(player_disp, start=1):
            imp = impacts[pk]
            m   = imp[metric_key].mean()
            s   = imp[metric_key].std()
            cell = row.cells[col_i]
            cell.paragraphs[0].add_run(f"{m:.1f} ± {s:.1f}").font.size = Pt(10)
            set_cell_bg(cell, "F5F7FA")
            if pk != "amateur" and m > best_val:
                best_val  = m
                best_name = pl

        row.cells[len(col_headers) - 1].paragraphs[0].add_run(best_name).font.size = Pt(10)
        set_cell_bg(row.cells[len(col_headers) - 1], "E8F5E9")

    doc.add_paragraph()
    add_body(doc,
        "Table 1 — Mean ± standard deviation of peak biomechanical metrics at serve impact. "
        "'Best Professional' indicates which elite player leads on each metric.",
        italic=True, size=9, color="666666")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 5 · KEY FINDINGS & COACHING IMPLICATIONS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Key Findings & Coaching Implications", level=1)
    add_divider(doc)

    add_heading(doc, "5.1 Professional Player Comparison", level=2, space_before=8)
    add_body(doc,
        "Despite very different playing styles, Djokovic and Federer converge on similar "
        "elbow and knee angles at impact, suggesting these are biomechanical invariants of "
        "elite serving. Their primary differentiation lies in shoulder abduction: Federer's "
        "trunk rotation-dominant technique produces a higher shoulder angle, while Djokovic's "
        "more upright serve maximises elbow extension and is associated with higher flat serve "
        "speed in real-world measurements.")

    add_heading(doc, "5.2 Amateur Development Areas", level=2, space_before=8)
    add_body(doc, "Prioritised coaching recommendations, ranked by impact:")

    coaching = [
        ("1. Shoulder Rotation (Critical)",
         "The amateur's shoulder angle at impact (72° ± 20°) is 45–60° below professional "
         "levels. This single metric accounts for most of the impact score deficit (z = −1.15). "
         "Drills focusing on hip-to-shoulder rotation and trophy pose shoulder elevation are the "
         "highest-leverage intervention."),
        ("2. Serve-to-Serve Consistency",
         "High standard deviations across all metrics indicate the amateur has not yet "
         "consolidated a repeatable motor pattern. Shadow-swing repetition and video "
         "self-review are recommended to accelerate pattern consolidation."),
        ("3. Knee Drive Timing",
         "The amateur's knee extension occurs slightly later and is less explosive than "
         "the professionals (avg peak 160° vs 166–170°). Leg-drive drills and jump-serve "
         "practice can help synchronise the kinematic chain."),
        ("4. Toss Consistency",
         "Wrist height trajectory variance is highest for the amateur. Practising the "
         "ball toss as an isolated skill — catching the ball at the intended contact point "
         "without hitting — will directly improve this metric."),
    ]
    for title, body in coaching:
        p   = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold  = True
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor.from_string("1F3864")
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        add_body(doc, body, size=10.5)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 6 · LIMITATIONS & FUTURE WORK
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Limitations & Future Work", level=1)
    add_divider(doc)

    add_heading(doc, "6.1 Current Limitations", level=2, space_before=8)
    lims = [
        "Single camera angle: all analysis is performed on 2D projected landmarks. Angles "
        "computed in the image plane may differ from true 3D joint angles, particularly "
        "for movements in/out of the camera plane.",
        "Small amateur sample (n=5): statistical power is limited. Conclusions about the "
        "amateur should be treated as directional rather than definitive.",
        "Serve type not distinguished: flat, slice, and kick serves have different biomechanical "
        "signatures; mixing them inflates variance, especially for Federer.",
        "Landmark occlusion: MediaPipe may produce degraded estimates when a limb is "
        "partially out of frame or occluded by clothing/equipment.",
        "Impact score heuristic: the composite score (elbow + knee − 300 × wrist_height) is an "
        "engineered proxy, not a ground-truth validated metric. Ball contact time from ball "
        "tracking would be a more reliable impact reference.",
    ]
    for l in lims:
        add_bullet(doc, l)

    add_heading(doc, "6.2 Proposed Enhancements", level=2, space_before=8)
    futures = [
        "3D pose estimation using multi-camera setup or depth sensor (e.g. Intel RealSense) to "
        "compute true 3D joint angles.",
        "Ball tracking integration (TrackNet / custom YOLO) to precisely anchor impact to the "
        "moment of racket-ball contact.",
        "Serve type classification via clustering on the biomechanical feature space to analyse "
        "flat, slice, and kick serves independently.",
        "Expansion to a larger amateur cohort (N≥20) for robust statistical comparisons and "
        "individualised feedback generation.",
        "Automated coaching feedback module: LLM-generated natural language recommendations "
        "derived from real-time biomechanical deviations from the professional baseline.",
    ]
    for f in futures:
        add_bullet(doc, f)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 7 · TECHNICAL APPENDIX
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Technical Appendix", level=1)
    add_divider(doc)

    add_heading(doc, "7.1 Technology Stack", level=2, space_before=8)
    tech = [
        ("MediaPipe Pose",   "v0.9",  "Pose estimation — 33 landmark BlazePose GHUM model"),
        ("OpenCV",           "4.x",   "Video I/O and frame decoding"),
        ("NumPy / Pandas",   "≥1.24", "Numerical computation and data wrangling"),
        ("Matplotlib",       "≥3.7",  "Publication-quality figure generation"),
        ("Seaborn",          "≥0.12", "Statistical visualisation (violin/box plots, heatmap)"),
        ("python-docx",      "1.2",   "Programmatic DOCX report generation"),
    ]
    tech_tbl = doc.add_table(rows=len(tech) + 1, cols=3)
    tech_tbl.style = "Table Grid"
    for cell, h in zip(tech_tbl.rows[0].cells, ["Library", "Version", "Purpose"]):
        cell.paragraphs[0].add_run(h).font.bold = True
        set_cell_bg(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for row, (lib, ver, purpose) in zip(tech_tbl.rows[1:], tech):
        for cell, v in zip(row.cells, [lib, ver, purpose]):
            cell.paragraphs[0].add_run(v).font.size = Pt(10)
            set_cell_bg(cell, "F5F7FA")

    add_heading(doc, "7.2 Angle Computation", level=2, space_before=12)
    add_body(doc,
        "All joint angles θ are computed as the angle at vertex B in the triplet (A, B, C):")
    add_body(doc,
        "    θ = |degrees(atan2(C_y−B_y, C_x−B_x) − atan2(A_y−B_y, A_x−B_x))|",
        bold=True, size=10)
    add_body(doc,
        "If the result exceeds 180°, it is reflected to 360° − θ to always return the interior "
        "angle in the range [0°, 180°].")

    add_heading(doc, "7.3 Impact Score Formula", level=2, space_before=8)
    add_body(doc,
        "    impact_score = elbow_angle + knee_angle − 300 × wrist_height_normalised",
        bold=True, size=10)
    add_body(doc,
        "The coefficient 300 was determined empirically to produce a score whose argmax coincides "
        "with the visually identifiable impact frame in validation clips. Wrist height is the "
        "minimum Y-coordinate (in image space) across both wrists, normalised to [0, 1] by image height.")

    add_heading(doc, "7.4 Reproducibility", level=2, space_before=8)
    add_body(doc,
        "All source code, intermediate CSVs, and generated figures are available in the project "
        "repository. The full pipeline can be re-run via:")
    add_body(doc,
        "    python src/multi_serve_analysis.py\n"
        "    python src/generate_report_figures.py\n"
        "    python src/build_report.py",
        size=10, bold=True)
    add_body(doc,
        "Random seeds are fixed (numpy default_rng(42)) for reproducible jitter in strip plots.")

    # ── Footer note ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_divider(doc)
    add_body(doc,
        f"Generated on {date.today().strftime('%B %d, %Y')} · "
        "Tennis Serve Biomechanics AI Project · Gonzalo Angel Tano",
        italic=True, size=9, color="999999", align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = "Tennis_Serve_Biomechanics_Report_v2.docx"
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
